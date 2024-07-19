import librosa.display
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO

def plot_spectrogram(input_path):
    # Đọc file âm thanh
    y, sr = librosa.load(input_path)

    # Tính toán biểu đồ tần số theo thời gian
    D = librosa.amplitude_to_db(librosa.stft(y), ref=np.max)

    # Vẽ biểu đồ
    plt.figure(figsize=(10, 4))
    librosa.display.specshow(D, sr=sr, x_axis='time', y_axis='log')
    plt.colorbar(format='%+2.0f dB')
    plt.title('Spectrogram')

    # Chuyển đổi biểu đồ thành mảng bytes
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')
    plt.close()

    return image_stream.getvalue()

def save_to_word(input_audio_path, output_word_path):
    # Tạo và lưu tệp Word
    doc = Document()
    doc.add_heading('Spectrogram of Audio File', level=1)

    # Vẽ biểu đồ và thêm vào tệp Word
    image_bytes = plot_spectrogram(input_audio_path)
    doc.add_picture(BytesIO(image_bytes), width=doc.default_width // 2)

    # Lưu tệp Word
    doc.save(output_word_path)

# Thay đổi đường dẫn file âm thanh đầu vào và đường dẫn đến nơi bạn muốn lưu tệp Word
input_audio_path = r'C:\Users\ducnh\Music\lan1.1.wav'
output_word_path = r'C:\Users\ducnh\Music\ducYeuNgoc.docx'

# Gọi hàm để vẽ biểu đồ và lưu vào tệp Word
save_to_word(input_audio_path, output_word_path)
